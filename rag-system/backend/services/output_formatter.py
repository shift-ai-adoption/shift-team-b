"""Output formatting: Markdown / Word (.docx) / Excel (.xlsx)."""
import io
import re
from datetime import datetime


def to_markdown(answer: str, query: str) -> bytes:
    md = (f"# 検索結果レポート\n\n**質問**: {query}\n\n"
          f"**作成日時**: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n---\n\n"
          f"{answer}\n")
    return md.encode("utf-8")


def to_word(answer: str, query: str) -> bytes:
    import docx
    doc = docx.Document()
    doc.add_heading("検索結果レポート", level=1)
    doc.add_paragraph(f"質問: {query}")
    doc.add_paragraph(f"作成日時: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    for line in answer.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        m = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if m:
            doc.add_heading(m.group(2), level=min(len(m.group(1)) + 1, 4))
        elif stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_excel(answer: str, query: str) -> bytes:
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "検索結果"
    ws.append(["質問", query])
    ws.append(["作成日時", datetime.now().strftime("%Y-%m-%d %H:%M")])
    ws.append([])
    # Parse markdown-like tables ("a | b | c") into rows, else dump lines
    for line in answer.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if "|" in stripped:
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                continue  # markdown separator row
            ws.append(cells)
        else:
            ws.append([stripped])
    for col in ("A", "B", "C"):
        ws.column_dimensions[col].width = 40
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


FORMATTERS = {"markdown": to_markdown, "word": to_word, "excel": to_excel}
EXT = {"markdown": "md", "word": "docx", "excel": "xlsx"}
MIME = {
    "markdown": "text/markdown; charset=utf-8",
    "word": ("application/vnd.openxmlformats-officedocument"
             ".wordprocessingml.document"),
    "excel": ("application/vnd.openxmlformats-officedocument"
              ".spreadsheetml.sheet"),
}
